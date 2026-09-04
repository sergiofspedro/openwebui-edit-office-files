"""Edit Office Files — Comments, tracked changes, and revision management."""

import io
import json
import os
import platform
import re
import sqlite3
import sys
import traceback
from copy import copy
from typing import Any, Dict, Optional

from .constants import *
from .utils import *

async def tracked_change(self, file_id: str, change_type: str, content: str, author: str = "Reviewer", paragraph_index: int = -1, output_filename: str = "", __user__=None, __request__=None) -> str:
    """Apply tracked changes (redlines) to a Word document with custom author name.
        Preserves original formatting and capitalization of the existing document.
    
        change_type: replace (use old_text|||new_text), insert (append text with redline), delete (mark paragraph for deletion)
        author: Name shown in Word's Track Changes (e.g., "Sergio Pedro")
        """
    try:
        import sqlite3 as s3
        conn2 = s3.connect(_DB_PATH)
        try:
            row = conn2.execute("SELECT filename, meta FROM file WHERE id=?", (file_id,)).fetchone()
            if not row:
                row = conn2.execute("SELECT filename, meta FROM file WHERE filename LIKE ?", (f"%{file_id}%",)).fetchone()
            if not row:
                return json.dumps({"error": "File not found"})
            filename = row[0]
            meta = json.loads(row[1]) if row[1] else {}
            data = _read_file_bytes(meta.get("path", file_id))
            if not data:
                return json.dumps({"error": "File not found on disk"})
        finally:
            conn2.close()

        # Check file type — track changes only supported for DOCX
        ftype = _detect_type(filename)
        if ftype != "docx":
            return json.dumps({"error": f"Track changes are only supported for DOCX files. {ftype.upper()} format does not support revision marks."})

        from docx import Document
        from docx_revisions import RevisionParagraph
        doc = Document(io.BytesIO(data))
        out_name = output_filename or filename
        results = []

        if change_type == "replace":
            parts = content.split("|||", 1)
            if len(parts) != 2:
                return json.dumps({"error": "Format: old_text|||new_text"})
            find_t, replace_t = parts
            for i, p in enumerate(doc.paragraphs):
                if paragraph_index >= 0 and i != paragraph_index:
                    continue
                if find_t in p.text:
                    rp = RevisionParagraph.from_paragraph(p)
                    cnt = rp.replace_tracked(find_t, replace_t, author=author)
                    results.append(f"Para {i}: {cnt} replacements")
        elif change_type == "insert":
            p = doc.add_paragraph()
            rp = RevisionParagraph.from_paragraph(p)
            rp.add_tracked_insertion(content, author=author)
            results.append("Inserted tracked text")
        elif change_type == "delete":
            idx = int(content) if content.isdigit() else paragraph_index
            if idx >= 0 and idx < len(doc.paragraphs):
                p = doc.paragraphs[idx]
                rp = RevisionParagraph.from_paragraph(p)
                rp.add_tracked_deletion(0, len(p.text), author=author)
                results.append(f"Marked para {idx} for deletion")

        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        url, name = await self._save_and_link(out.read(), out_name, __request__)
        if url:
            return f"[{name}]({url})\n\nTracked changes by '{author}':\n" + "\n".join(results)
        return json.dumps({"error": "Could not save file"})
    except Exception as e:
        return json.dumps({"error": str(e), "traceback": traceback.format_exc()})


async def manage_revisions(self, file_id: str, action: str, output_filename: str = "", __user__=None, __request__=None) -> str:
    """List, accept_all or reject_all tracked changes in a Word document."""
    try:
        import sqlite3 as s3
        conn2 = s3.connect(_DB_PATH)
        try:
            row = conn2.execute("SELECT filename, meta FROM file WHERE id=?", (file_id,)).fetchone()
            if not row:
                row = conn2.execute("SELECT filename, meta FROM file WHERE filename LIKE ?", (f"%{file_id}%",)).fetchone()
            if not row:
                return json.dumps({"error": "File not found"})
            filename = row[0]
            meta = json.loads(row[1]) if row[1] else {}
            data = _read_file_bytes(meta.get("path", file_id))
            if not data:
                return json.dumps({"error": "File not found on disk"})
        finally:
            conn2.close()

        # Check file type — revision management only supported for DOCX
        ftype = _detect_type(filename)
        if ftype != "docx":
            return json.dumps({"error": f"Revision management is only supported for DOCX files. {ftype.upper()} format does not support revision marks."})

        from docx_revisions import RevisionDocument

        if action == "list":
            rdoc = RevisionDocument(io.BytesIO(data))
            revs = []
            for para in rdoc.paragraphs:
                try:
                    rp = RevisionParagraph.from_paragraph(para)
                    if rp.has_track_changes:
                        for ins in rp.insertions:
                            revs.append({"type": "insertion", "author": ins.author, "text": ins.text[:100]})
                        for d in rp.deletions:
                            revs.append({"type": "deletion", "author": d.author, "text": d.text[:100]})
                except Exception:
                    pass
            return json.dumps({"revisions": revs, "count": len(revs)}, indent=2)

        out_name = output_filename or filename
        rdoc = RevisionDocument(io.BytesIO(data))
        if action == "accept_all":
            rdoc.accept_all()
            msg = "All track changes accepted"
        elif action == "reject_all":
            rdoc.reject_all()
            msg = "All track changes rejected"
        else:
            return json.dumps({"error": f"Unknown action: {action}"})

        out = io.BytesIO()
        rdoc.save(out)
        out.seek(0)
        url, name = await self._save_and_link(out.read(), out_name, __request__)
        if url:
            return f"[{name}]({url})\n\n{msg}."
        return json.dumps({"error": "Could not save file"})
    except Exception as e:
        return json.dumps({"error": str(e), "traceback": traceback.format_exc()})



# --- v3.2.0: ODF Write ---


async def create_odf(self, content: str, filename: str = "document", format: str = "odt", raw_text: bool = False, __user__=None, __request__=None) -> str:
    """Create a new ODF file (.odt, .ods, .odp).

        Args:
            content: Content for the file (CSV for .ods, text for .odt, markdown for .odp)
            filename: Output filename
            format: 'odt', 'ods', or 'odp'
            raw_text: If True, skip text formatting
        """
    from odf.opendocument import OpenDocumentText, OpenDocumentSpreadsheet, OpenDocumentPresentation
    from odf.text import P, H
    from odf.table import Table, TableRow, TableCell
    import io
    fmt_mode = "preserve" if raw_text else "format"
    
    try:
        if format == "ods":
            doc = OpenDocumentSpreadsheet()
            lines = content.split('\n')
            table = Table(name="Sheet1")
            for line in lines:
                line = line.strip()
                if not line: continue
                cells = [c.strip() for c in line.split(',')]
                row = TableRow()
                for c in cells:
                    cell = TableCell()
                    cell.addElement(P(text=_format_text(c, mode=fmt_mode)))
                    row.addElement(cell)
                table.addElement(row)
            doc.spreadsheet.addElement(table)
        elif format == "odp":
            doc = OpenDocumentPresentation()
            for line in content.split('\n'):
                line = line.strip()
                if not line: continue
                if line.startswith('# '):
                    doc.presentation.addElement(H(outlinelevel=1, text=line[2:]))
                else:
                    doc.presentation.addElement(P(text=_format_text(line, mode=fmt_mode)))
        else:
            doc = OpenDocumentText()
            for line in content.split('\n'):
                line = line.strip()
                if not line: continue
                if line.startswith('# '):
                    doc.text.addElement(H(outlinelevel=1, text=line[2:]))
                elif line.startswith('## '):
                    doc.text.addElement(H(outlinelevel=2, text=line[3:]))
                else:
                    doc.text.addElement(P(text=_format_text(line, mode=fmt_mode)))
        
        buf = io.BytesIO()
        doc.write(buf)
        buf.seek(0)
        url, fname = await self._save_and_link(buf.getvalue(), f"{filename}.{format}", __request__, __user__=__user__)
        if url:
            return f"ODF file created: [{fname}]({url})"
        return json.dumps({"error": "Could not save file"})
    except ImportError:
        return json.dumps({"error": "odfpy not installed. Install with: pip install odfpy"})
    except Exception as e:
        return json.dumps({"error": str(e)})

# --- v3.2.0: Format Conversion ---


async def add_comment(self, file_id: str, text: str, author: str = "Reviewer", paragraph_index: int = 0, cell_ref: str = "A1", slide_num: int = 1, page_num: int = 1, __user__=None, __request__=None) -> str:
    """Add a review comment to a Word, Excel, PowerPoint, or PDF file.

        Args:
            file_id: File ID to comment on
            text: Comment text
            author: Name shown in the comment (e.g., "Sergio Pedro")
            paragraph_index: For DOCX: paragraph index to attach comment to (default 0)
            cell_ref: For XLSX: cell reference (default "A1")
            slide_num: For PPTX: slide number (default 1)
            page_num: For PDF: page number to attach comment to (default 1)
        """
    import io
    file_bytes, filename, ftype = self._resolve_file(file_id)
    if not file_bytes: return json.dumps({"error": "File not found"})

    if ftype == "docx":
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))

        # Ensure we have a paragraph to comment on
        if not doc.paragraphs:
            doc.add_paragraph("")
        if paragraph_index >= len(doc.paragraphs):
            paragraph_index = 0
        para = doc.paragraphs[paragraph_index]

        comment = doc.add_comment(text, author=author)
        para.add_comment(comment)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        url, fname = await self._save_and_link(buf.getvalue(), filename, __request__)
        if url: return f"Comment added by {author} on paragraph {paragraph_index}: [{fname}]({url})"
        return json.dumps({"error": "Could not save file"})

    elif ftype == "xlsx":
        from openpyxl import load_workbook
        from openpyxl.comments import Comment
        wb = load_workbook(io.BytesIO(file_bytes))
        ws = wb.active
        comment = Comment(text, author)
        ws[cell_ref].comment = comment
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        url, fname = await self._save_and_link(buf.getvalue(), filename, __request__)
        if url: return f"Comment added by {author} on cell {cell_ref}: [{fname}]({url})"
        return json.dumps({"error": "Could not save file"})

    elif ftype == "pptx":
        try:
            import zipfile
            import uuid as _uuid
            from datetime import datetime, timezone
            from lxml import etree


            slide_name = "ppt/slides/slide%d.xml" % slide_num
            rels_name = "ppt/slides/_rels/slide%d.xml.rels" % slide_num
            modern_name = "ppt/comments/commentModern%d.xml" % slide_num
            authors_name = "ppt/commentsAuthors.xml"

            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                entries = {name: z.read(name) for name in z.namelist()}

            if slide_name not in entries:
                return json.dumps({"error": "Slide %d not found in presentation" % slide_num})

            etree.register_namespace("p", _P_NS)
            etree.register_namespace("p14", _P14_NS)
            etree.register_namespace("r", _R_NS)

            utc_now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.000")
            para_guid = "{" + str(_uuid.uuid4()).upper() + "}"

            # --- modern comment part: compute next comment idx ---
            if modern_name in entries:
                mroot = _safe_etree_fromstring(entries[modern_name], context=modern_name)
                if mroot is None:
                    # Could not parse existing modern comments — start fresh
                    mroot = etree.Element("{%s}cmLst" % _P_NS)
                max_idx = 0
                for el in mroot:
                    if el.tag == "{%s}cm" % _P_NS:
                        try:
                            max_idx = max(max_idx, int(el.get("idx") or 0))
                        except (TypeError, ValueError):
                            pass
                next_idx = max_idx + 1
            else:
                mroot = etree.Element("{%s}cmLst" % _P_NS)
                next_idx = 1

            # --- commentsAuthors part: resolve or create author id ---
            if authors_name in entries:
                aroot = _safe_etree_fromstring(entries[authors_name], context=authors_name)
                if aroot is None:
                    aroot = etree.Element("{%s}cmAuthorLst" % _P14_NS)
            else:
                aroot = etree.Element("{%s}cmAuthorLst" % _P14_NS)

            author_id = None
            existing_ids = []
            author_els = []
            for el in aroot:
                if el.tag == "{%s}cmAuthor" % _P14_NS:
                    author_els.append(el)
                    try:
                        existing_ids.append(int(el.get("id") or 0))
                    except (TypeError, ValueError):
                        existing_ids.append(0)
                    if el.get("name") == author:
                        author_id = el.get("id")

            if author_id is not None:
                for el in author_els:
                    if el.get("name") == author:
                        el.set("lastIdx", str(next_idx))
                        break
            else:
                new_id = (max(existing_ids) + 1) if existing_ids else 0
                initials = "".join(w[0] for w in re.split(r"[\s._\-]+", author.strip()) if w)[:4].upper() or "AU"
                new_author = etree.SubElement(aroot, "{%s}cmAuthor" % _P14_NS)
                new_author.set("id", str(new_id))
                new_author.set("name", author)
                new_author.set("initials", initials)
                new_author.set("lastIdx", "1")
                new_author.set("clrIndex", str(len(author_els) % 14))
                author_id = str(new_id)

            # --- append modern comment ---
            cm = etree.SubElement(mroot, "{%s}cm" % _P_NS)
            cm.set("authorId", author_id)
            cm.set("dt", utc_now)
            cm.set("idx", str(next_idx))
            pos = etree.SubElement(cm, "{%s}pos" % _P_NS)
            pos.set("x", "100")
            pos.set("y", "100")
            cm_txt = etree.SubElement(cm, "{%s}text" % _P_NS)
            cm_txt.text = text
            cm_extLst = etree.SubElement(cm, "{%s}extLst" % _P_NS)
            cm_ext = etree.SubElement(cm_extLst, "{%s}ext" % _P_NS)
            cm_ext.set("uri", "{D6B160D4-4F5E-48AF-9E34-8E18A86A1F0A}")
            comment_ex = etree.SubElement(cm_ext, "{%s}commentEx" % _P14_NS)
            comment_ex.set("paraId", para_guid)
            comment_ex.set("dt", utc_now)
            comment_ex.set("parentIdx", "0")
            ce_txt = etree.SubElement(comment_ex, "{%s}text" % _P14_NS)
            ce_txt.text = text
            ce_author = etree.SubElement(comment_ex, "{%s}authorId" % _P14_NS)
            ce_author.text = author_id

            # --- [Content_Types].xml: add overrides if missing ---
            ct_root = _safe_etree_fromstring(entries["[Content_Types].xml"], context="[Content_Types].xml")
            if ct_root is None:
                # Rebuild minimal [Content_Types].xml — better than losing the
                # entire content-type table on a single bad XML read.
                ct_root = etree.Element("{%s}Types" % _CT_NS, nsmap={None: _CT_NS})
            has_authors_ct = any(
                el.tag == "{%s}Override" % _CT_NS and el.get("PartName") == "/ppt/commentsAuthors.xml"
                for el in ct_root
            )
            if not has_authors_ct:
                o = etree.SubElement(ct_root, "{%s}Override" % _CT_NS)
                o.set("PartName", "/ppt/commentsAuthors.xml")
                o.set("ContentType", _CT_AUTHORS)
            has_modern_ct = any(
                el.tag == "{%s}Override" % _CT_NS and el.get("PartName") == "/ppt/comments/commentModern%d.xml" % slide_num
                for el in ct_root
            )
            if not has_modern_ct:
                o = etree.SubElement(ct_root, "{%s}Override" % _CT_NS)
                o.set("PartName", "/ppt/comments/commentModern%d.xml" % slide_num)
                o.set("ContentType", _CT_MODERN)

            # --- slide rels: add commentsModern relationship ---
            if rels_name in entries:
                rroot = _safe_etree_fromstring(entries[rels_name], context=rels_name)
                if rroot is None:
                    rroot = etree.Element("{%s}Relationships" % _PKG_REL_NS, nsmap={None: _PKG_REL_NS})
            else:
                rroot = etree.Element("{%s}Relationships" % _PKG_REL_NS, nsmap={None: _PKG_REL_NS})
            max_rid = 0
            for el in rroot:
                rid = el.get("Id") or ""
                m = re.match(r"rId(\d+)$", rid)
                if m:
                    max_rid = max(max_rid, int(m.group(1)))
            new_rid = max_rid + 1
            target = "../comments/commentModern%d.xml" % slide_num
            # Deduplicate: reuse existing Relationship to same target (avoid duplicate on repeat calls)
            existing_rel = None
            for child in rroot:
                if child.tag == "{%s}Relationship" % _PKG_REL_NS and child.get("Type") == _CM_REL_TYPE and child.get("Target") == target:
                    existing_rel = child
                    break
            if existing_rel is not None:
                existing_rel.set("Id", "rId%d" % new_rid)
            else:
                rel_el = etree.SubElement(rroot, "{%s}Relationship" % _PKG_REL_NS)
                rel_el.set("Id", "rId%d" % new_rid)
                rel_el.set("Type", _CM_REL_TYPE)
                rel_el.set("Target", target)

            # --- slide XML: attach commentRel extension ---
            sroot = _safe_etree_fromstring(entries[slide_name], context=slide_name)
            if sroot is None:
                raise ValueError(
                    f"Could not parse slide XML '{slide_name}' — refusing to "
                    "overwrite a slide whose structure we cannot read."
                )
            extLst = None
            for child in sroot:
                if child.tag == "{%s}extLst" % _P_NS:
                    extLst = child
                    break
            if extLst is None:
                extLst = etree.SubElement(sroot, "{%s}extLst" % _P_NS)
            # Deduplicate: reuse existing commentRel ext if present (avoid duplicate URI on repeat calls)
            slide_ext = None
            for child in extLst:
                if child.tag == "{%s}ext" % _P_NS and child.get("uri") == "{6950BFC3-D8DA-4A85-94F7-54DA5524770B}":
                    slide_ext = child
                    break
            if slide_ext is None:
                slide_ext = etree.SubElement(extLst, "{%s}ext" % _P_NS)
                slide_ext.set("uri", "{6950BFC3-D8DA-4A85-94F7-54DA5524770B}")
            # Update or create commentRel child
            comment_rel = None
            for child in slide_ext:
                if child.tag == "{%s}commentRel" % _P14_NS:
                    comment_rel = child
                    break
            if comment_rel is None:
                comment_rel = etree.SubElement(slide_ext, "{%s}commentRel" % _P14_NS)
            comment_rel.set("{%s}id" % _R_NS, "rId%d" % new_rid)

            # --- serialize modified parts ---
            def _ser(root):
                return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

            entries[modern_name] = _ser(mroot)
            entries[authors_name] = _ser(aroot)
            entries[rels_name] = _ser(rroot)
            entries[slide_name] = _ser(sroot)
            entries["[Content_Types].xml"] = _ser(ct_root)

            # --- rebuild zip (OPC: [Content_Types].xml first) ---
            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
                zout.writestr("[Content_Types].xml", entries["[Content_Types].xml"])
                for name in entries:
                    if name != "[Content_Types].xml":
                        zout.writestr(name, entries[name])
            buf.seek(0)

            url, fname = await self._save_and_link(buf.getvalue(), filename, __request__)
            if url:
                return f"Comment added by {author} on slide {slide_num}: [{fname}]({url})"
            return json.dumps({"error": "Could not save file"})
        except Exception as e:
            return json.dumps({"error": str(e), "traceback": traceback.format_exc()})

    elif ftype == "pdf":
        try:
            import fitz
            pdf = fitz.open(stream=file_bytes, filetype="pdf")
            if page_num < 1 or page_num > pdf.page_count:
                total_pages = pdf.page_count
                pdf.close()
                return json.dumps({"error": f"Page {page_num} not found in PDF (has {total_pages} pages)."})
            page = pdf.load_page(page_num - 1)
            annot = page.add_text_annot(fitz.Point(72, 72), text)
            annot.set_info(title=author, content=text, subject="Comment")
            buf = io.BytesIO()
            pdf.save(buf)
            pdf.close()
            buf.seek(0)
            url, fname = await self._save_and_link(buf.getvalue(), filename, __request__)
            if url:
                return f"Comment added by {author} on page {page_num}: [{fname}]({url})"
            return json.dumps({"error": "Could not save file"})
        except Exception as e:
            return json.dumps({"error": str(e), "traceback": traceback.format_exc()})

    else:
        return json.dumps({"error": f"Comments not supported for {ftype}. Supported: DOCX, XLSX, PPTX, PDF."})


__all__ = [
    "add_comment",
    "create_odf",
    "manage_revisions",
    "tracked_change",
]
