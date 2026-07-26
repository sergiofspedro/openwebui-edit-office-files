    async def tracked_change(self, file_id: str, change_type: str, content: str, author: str = "Reviewer", paragraph_index: int = -1, output_filename: str = "", __user__=None, __request__=None) -> str:
        """Apply tracked changes (redlines) to a Word document with custom author name.
    
        change_type: replace (use old_text|||new_text), insert (append text with redline), delete (mark paragraph for deletion)
        author: Name shown in Word's Track Changes (e.g., "Sergio Pedro")
        """
        try:
            import sqlite3 as s3
            conn2 = s3.connect(_DB_PATH)
            row = conn2.execute("SELECT filename, meta FROM file WHERE id=?", (file_id,)).fetchone()
            if not row:
                row = conn2.execute("SELECT filename, meta FROM file WHERE filename LIKE ?", (f"%{file_id}%",)).fetchone()
            if not row:
                conn2.close()
                return json.dumps({"error": "File not found"})
            filename = row[0]
            meta = json.loads(row[1]) if row[1] else {}
            fp = meta.get("path", file_id)
            if not os.path.exists(fp):
                fp = os.path.join(_get_owui_data_dir(), "uploads", os.path.basename(fp))
            if not os.path.exists(fp):
                conn2.close()
                return json.dumps({"error": "File not found on disk"})
            with open(fp, "rb") as f:
                data = f.read()
            conn2.close()
    
            from docx import Document
            from docx_revisions import RevisionParagraph
            doc = Document(io.BytesIO(data))
            out_name = output_filename or filename
            results = []
    
            if change_type == "replace":
                parts = content.split("|||", 1)
                if len(parts) != 2:
                    return json.dumps({"error": "Format: old_text|||new_text"})
                find_t, replace_t = parts
                for i, p in enumerate(doc.paragraphs):
                    if paragraph_index >= 0 and i != paragraph_index:
                        continue
                    if find_t in p.text:
                        rp = RevisionParagraph.from_paragraph(p)
                        cnt = rp.replace_tracked(find_t, replace_t, author=author)
                        results.append(f"Para {i}: {cnt} replacements")
            elif change_type == "insert":
                p = doc.add_paragraph()
                rp = RevisionParagraph.from_paragraph(p)
                rp.add_tracked_insertion(content, author=author)
                results.append("Inserted tracked text")
            elif change_type == "delete":
                idx = int(content) if content.isdigit() else paragraph_index
                if idx >= 0 and idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]
                    rp = RevisionParagraph.from_paragraph(p)
                    rp.add_tracked_deletion(0, len(p.text), author=author)
                    results.append(f"Marked para {idx} for deletion")
    
            out = io.BytesIO()
            doc.save(out)
            out.seek(0)
            url, name = await self._save_and_link(out.read(), out_name, __request__)
            if url:
                return f"[{name}]({url})\n\nTracked changes by '{author}':\n" + "\n".join(results)
            return json.dumps({"error": "Could not save file"})
        except Exception as e:
            return json.dumps({"error": str(e), "traceback": traceback.format_exc()})


    async def manage_revisions(self, file_id: str, action: str, output_filename: str = "", __user__=None, __request__=None) -> str:
        """List, accept_all or reject_all tracked changes in a Word document."""
        try:
            import sqlite3 as s3
            conn2 = s3.connect(_DB_PATH)
            row = conn2.execute("SELECT filename, meta FROM file WHERE id=?", (file_id,)).fetchone()
            if not row:
                row = conn2.execute("SELECT filename, meta FROM file WHERE filename LIKE ?", (f"%{file_id}%",)).fetchone()
            if not row:
                conn2.close()
                return json.dumps({"error": "File not found"})
            filename = row[0]
            meta = json.loads(row[1]) if row[1] else {}
            fp = meta.get("path", file_id)
            if not os.path.exists(fp):
                fp = os.path.join(_get_owui_data_dir(), "uploads", os.path.basename(fp))
            with open(fp, "rb") as f:
                data = f.read()
            conn2.close()
    
            from docx_revisions import RevisionDocument
    
            if action == "list":
                rdoc = RevisionDocument(io.BytesIO(data))
                revs = []
                for para in rdoc.paragraphs:
                    try:
                        rp = RevisionParagraph.from_paragraph(para)
                        if rp.has_track_changes:
                            for ins in rp.insertions:
                                revs.append({"type": "insertion", "author": ins.author, "text": ins.text[:100]})
                            for d in rp.deletions:
                                revs.append({"type": "deletion", "author": d.author, "text": d.text[:100]})
                    except Exception:
                        pass
                return json.dumps({"revisions": revs, "count": len(revs)}, indent=2)
    
            out_name = output_filename or filename
            rdoc = RevisionDocument(io.BytesIO(data))
            if action == "accept_all":
                rdoc.accept_all()
                msg = "All track changes accepted"
            elif action == "reject_all":
                rdoc.reject_all()
                msg = "All track changes rejected"
            else:
                return json.dumps({"error": f"Unknown action: {action}"})
    
            out = io.BytesIO()
            rdoc.save(out)
            out.seek(0)
            url, name = await self._save_and_link(out.read(), out_name, __request__)
            if url:
                return f"[{name}]({url})\n\n{msg}."
            return json.dumps({"error": "Could not save file"})
        except Exception as e:
            return json.dumps({"error": str(e), "traceback": traceback.format_exc()})



    # --- v3.2.0: ODF Write ---
