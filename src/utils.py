"""Edit Office Files — shared utilities and formatting helpers."""
import os
import re
import sys
import sqlite3
import io
import base64 as _b64_mod
from typing import Optional
from .constants import _data_dir, _DB_PATH, _UPLOAD_DIR, _EXPORT_DIR

def _resolve_file_path(file_id: str) -> Optional[str]:
    """Resolve an Open WebUI file UUID to an absolute disk path."""
    try:
        conn = sqlite3.connect(f"file:{_DB_PATH}?mode=ro", uri=True)
        try:
            row = conn.execute(
                "SELECT path FROM file WHERE id = ?", (file_id,)
            ).fetchone()
        finally:
            conn.close()
    except Exception as exc:
        print(f"[office] DB lookup failed for {file_id}: {exc}", file=sys.stderr)
        return None

    if not row or not row[0]:
        # Fallback: try by filename
        try:
            conn = sqlite3.connect(f"file:{_DB_PATH}?mode=ro", uri=True)
            try:
                row = conn.execute(
                    "SELECT path FROM file WHERE filename LIKE ?",
                    (f"%{file_id}%",),
                ).fetchone()
            finally:
                conn.close()
        except Exception:
            pass

    if not row or not row[0]:
        print(f"[office] No path for file_id {file_id}", file=sys.stderr)
        return None

    path = row[0]
    if os.path.isfile(path):
        return path

    # Fallback: uploads directory
    candidate = os.path.join(_UPLOAD_DIR, os.path.basename(path))
    if os.path.isfile(candidate):
        return candidate

    # Last resort: UUID prefix match in uploads
    prefix = file_id.split("-")[0] if "-" in file_id else file_id[:8]
    if os.path.isdir(_UPLOAD_DIR):
        for name in os.listdir(_UPLOAD_DIR):
            if name.startswith(prefix):
                candidate = os.path.join(_UPLOAD_DIR, name)
                if os.path.isfile(candidate):
                    return candidate

    print(f"[office] File not found on disk: {path}", file=sys.stderr)
    return None


def _read_file_bytes(file_id: str) -> Optional[bytes]:
    """Return raw bytes for an Open WebUI file."""
    path = _resolve_file_path(file_id)
    if path is None:
        return None
    # Path traversal guard: ensure resolved path stays inside allowed directories
    try:
        _abs = os.path.realpath(path)
        _allowed = False
        for _base in (_UPLOAD_DIR, _EXPORT_DIR,
                       os.path.join(os.environ.get("OPEN_WEBUI_DATA_DIR", ""), "data"),
                       _get_owui_data_dir()):
            if _base and _abs.startswith(os.path.realpath(_base) + os.sep):
                _allowed = True
                break
        if not _allowed:
            print(f"[office] Path traversal blocked: {path}", file=sys.stderr)
            return None
    except Exception as exc:
        print(f"[office] Path traversal check failed for {path}: {exc}", file=sys.stderr)
        return None
    try:
        with open(path, "rb") as fh:
            return fh.read()
    except Exception as exc:
        print(f"[office] Read failed for {path}: {exc}", file=sys.stderr)
        return None


def _detect_type(filename: str) -> str:
    """Detect file type from extension."""
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".xlsx",):
        return "xlsx"
    if ext in (".xls",):
        return "xls"
    if ext in (".docx",):
        return "docx"
    if ext in (".pptx",):
        return "pptx"
    if ext in (".ppt",):
        return "ppt"
    if ext in (".odt",):
        return "odt"
    if ext in (".ods",):
        return "ods"
    if ext in (".odp",):
        return "odp"
    if ext in (".csv",):
        return "csv"
    if ext in (".pdf",):
        return "pdf"
    return "unknown"


def _cell_value(cell):
    """Extract a JSON-safe value from an openpyxl cell."""
    import datetime

    v = cell.value
    if v is None:
        return None
    if isinstance(v, (datetime.datetime, datetime.date)):
        return v.isoformat()
    if isinstance(v, (int, float, str, bool)):
        return v
    return str(v)


def _xls_to_xlsx(xls_data: bytes) -> bytes:
    """Convert .xls bytes to .xlsx bytes using xlrd + openpyxl.

    Returns an in-memory .xlsx workbook as bytes.
    """
    import xlrd
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    xls_book = xlrd.open_workbook(file_contents=xls_data)
    xlsx_wb = openpyxl.Workbook()
    # Remove the default sheet; we'll add one per xls sheet
    xlsx_wb.remove(xlsx_wb.active)

    for sheet_idx in range(xls_book.nsheets):
        xls_sheet = xls_book.sheet_by_index(sheet_idx)
        ws = xlsx_wb.create_sheet(title=xls_sheet.name[:31])  # Excel 31-char limit

        for rx in range(xls_sheet.nrows):
            for cx in range(xls_sheet.ncols):
                cell = xls_sheet.cell(rx, cx)
                value = cell.value

                # xlrd date handling: if cell type is XL_CELL_DATE, convert
                if cell.ctype == xlrd.XL_CELL_DATE:
                    try:
                        import datetime as _dt
                        value = _dt.datetime(*xlrd.xldate_as_tuple(value, xls_book.datemode))
                    except Exception:
                        pass
                elif cell.ctype == xlrd.XL_CELL_BOOLEAN:
                    value = bool(value)

                ws.cell(row=rx + 1, column=cx + 1, value=value)

        # Auto-fit column widths (rough estimate)
        for col_cells in ws.columns:
            max_len = 0
            for cell in col_cells:
                try:
                    max_len = max(max_len, len(str(cell.value or "")))
                except Exception:
                    pass
            letter = openpyxl.utils.get_column_letter(col_cells[0].column)
            ws.column_dimensions[letter].width = min(max_len + 2, 50)

    out = io.BytesIO()
    xlsx_wb.save(out)
    xlsx_wb.close()
    xls_book.release_resources()
    out.seek(0)
    return out.read()


def _format_text(text: str, mode: str = "format") -> str:
    """Apply consistent text formatting: normalize dashes, sentence case, preserve acronyms.

    Args:
        text: The text to format.
        mode: "format" (default) applies sentence case + acronym preservation + dash normalization.
              "preserve" returns text unchanged.
    """
    if mode == "preserve":
        return text
    if not isinstance(text, str) or not text or text.startswith('='):
        return text
    
    # 1. Normalize dashes
    text = text.replace('\u2014', '-').replace('\u2015', '-').replace('\u2013', '-')
    
    # 2. Split into sentences and apply sentence case
    acronyms = {'API', 'PDF', 'HTML', 'CSS', 'JSON', 'SQL', 'AI', 'UI', 'UX', 'ID', 'URL', 'HTTP', 'HTTPS', 'FTP', 'SSH', 'DNS', 'IP', 'TCP', 'UDP', 'SSL', 'TLS', 'REST', 'CRUD', 'YAML', 'XML', 'SVG', 'PNG', 'JPG', 'GIF', 'CSV', 'DOCX', 'PPTX', 'XLSX', 'DOC', 'PPT', 'XLS', 'ISO', 'RGB', 'CMYK', 'PHP', 'NPM', 'YARN', 'CLI', 'GUI', 'IDE', 'SDK', 'JDK', 'JRE', 'JVM', 'DB', 'SQLite', 'MySQL', 'NoSQL', 'OAuth', 'JWT', 'CORS', 'MVC', 'MVP', 'MVVM', 'SPA', 'PWA', 'SSR', 'CSR', 'SSG', 'ISR', 'CDN', 'DHCP', 'NAT', 'VPN', 'LAN', 'WAN', 'MAC', 'BIOS', 'UEFI', 'GPT', 'MBR', 'RAID', 'NAS', 'SAN', 'AWS', 'GCP', 'Azure', 'SaaS', 'PaaS', 'IaaS', 'FaaS', 'CaaS', 'K8s', 'Docker', 'Kubernetes'}
    
    import re as _re
    
    # Protect acronyms with placeholders (longest first to avoid partial matches)
    acronym_map = {}
    for i, acro in enumerate(sorted(acronyms, key=len, reverse=True)):
        placeholder = f'\x00{i:04d}\x00'
        pattern = _re.compile(r'(?<![a-zA-Z])' + _re.escape(acro) + r'(?![a-zA-Z])', _re.IGNORECASE)
        text = pattern.sub(placeholder, text)
        acronym_map[placeholder] = acro
    
    # Smart sentence splitting: avoid splitting on abbreviations
    # Heuristic 1: A word containing an internal period (e.g., "U.S.A.", "e.g.", "i.e.", "Ph.D.")
    #              is an abbreviation — don't split after it (handles U.S.A., U.K., E.U., a.k.a., w.r.t., etc.)
    # Mini-list: short common abbreviations without internal periods (Dr., Mr., St., etc.)
    # Note: month abbreviations (Jan, Feb, etc.) are included since they're commonly followed by
    # a number (e.g., "Jan 15") which would pass the sentence-split check.
    _MINI_ABBREV = {'Dr', 'Mr', 'Mrs', 'Ms', 'Prof', 'Sr', 'Jr', 'St', 'vs', 'etc',
                    'Jan', 'Feb', 'Mar', 'Apr', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec'}

    _abbrev_pattern = '|'.join(_re.escape(a) for a in sorted(_MINI_ABBREV, key=len, reverse=True))
    _sentence_re = _re.compile(
        r'(?<!\b(?:' + _abbrev_pattern + r'))'   # not after mini-abbrev
        r'(?<!\.[A-Za-z])'                         # not after word with internal period
        r'(?<=[.!?])\s+'                            # after .!? + space
    )
    sentences = _sentence_re.split(text)

    # Sentence case with intentional capitalization preservation
    # Rule: words with ANY uppercase letter are preserved as-is (proper nouns,
    # acronyms not in the list, product names, the pronoun "I", etc.).
    # Words that are entirely lowercase get sentence-cased.
    formatted_sentences = []
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        
        # Split into words while preserving spacing
        words = s.split(' ')
        formatted_words = []
        
        for word in words:
            if not word:
                formatted_words.append(word)
                continue
            
            # Skip placeholder-protected acronyms (they start with \x00)
            if '\x00' in word:
                formatted_words.append(word)
                continue
            
            # Check if word has intentional capitalization
            # (any uppercase letter means the user intentionally capitalized it)
            has_uppercase = any(c.isupper() for c in word)
            
            if has_uppercase:
                # Preserve original capitalization
                formatted_words.append(word)
            else:
                # Safe to lowercase
                formatted_words.append(word.lower())
        
        s = ' '.join(formatted_words)
        
        # Capitalize first letter of the first word in the sentence
        if s and s[0].isalpha():
            s = s[0].upper() + s[1:]
        
        formatted_sentences.append(s)
    
    text = ' '.join(formatted_sentences)
    
    # Restore acronyms
    for placeholder, acro in acronym_map.items():
        text = text.replace(placeholder, acro)
    
    return text


def _parse_inline_md(text: str) -> list[tuple[str, dict]]:
    """Parse inline markdown in *text* and return segments with formatting metadata.

    Returns a list of ``(plain_text, formatting_dict)`` tuples. The formatting dict
    may contain the keys ``bold``, ``italic``, ``code``, or ``link``.

    Parsing order (priority, first matched wins):
        1. Links — ``[text](url)``
        2. Bold — ``**text**``
        3. Italic — ``*text*``  (conflicts with ``* `` bullet lists are avoided)
        4. Inline code — `` `code` ``
    """
    import re as _re

    if not isinstance(text, str):
        text = str(text) if text is not None else ""

    if text == "":
        return [("", {})]

    # Regexes — link first to prevent its brackets from being consumed by bold/italic
    _link_re = _re.compile(r'\[(.+?)\]\(([^()]*(?:\([^()]*\)[^()]*)*)\)')
    _bold_re = _re.compile(r'\*\*(.+?)\*\*')
    _italic_re = _re.compile(r'(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)')
    _code_re = _re.compile(r'`(.+?)`')

    segments: list[tuple[str, dict]] = []

    # Work on a mutable list of (start, end, kind, data) markers.
    # Kind: 'link', 'bold', 'italic', 'code'
    markers: list[tuple[int, int, str, str]] = []

    # Links
    for m in _link_re.finditer(text):
        markers.append((m.start(), m.end(), 'link', (m.group(1), m.group(2))))  # (label, url)

    # Bold
    for m in _bold_re.finditer(text):
        # Only add if this range doesn't overlap with a higher-priority marker
        markers.append((m.start(), m.end(), 'bold', m.group(1)))

    # Italic
    for m in _italic_re.finditer(text):
        markers.append((m.start(), m.end(), 'italic', m.group(1)))

    # Code
    for m in _code_re.finditer(text):
        markers.append((m.start(), m.end(), 'code', m.group(1)))

    # Sort by start position, then by end position (longer span first wins ties)
    markers.sort(key=lambda x: (x[0], -x[1]))

    # Remove overlapping markers: keep the earliest-starting, and if tied,
    # the longer one; discard any later marker whose start < retained end.
    cleaned: list[tuple[int, int, str, str]] = []
    for mkr in markers:
        if not cleaned or mkr[0] >= cleaned[-1][1]:
            cleaned.append(mkr)
        else:
            # Overlap — keep only if this marker starts at the same position
            # as the last kept one but is longer (shouldn't happen with our
            # sort, but be safe) or is a link that overlaps with something else.
            pass

    markers = cleaned

    pos = 0
    for start, end, kind, data in markers:
        # Plain text before this marker
        if start > pos:
            segments.append((text[pos:start], {}))
        # The formatted segment
        fmt: dict = {}
        if kind == 'link':
            label, url = data
            fmt['link'] = url
            text_part = label
        else:
            if kind == 'bold':
                fmt['bold'] = True
            elif kind == 'italic':
                fmt['italic'] = True
            elif kind == 'code':
                fmt['code'] = True
            text_part = data
        segments.append((text_part, fmt))
        pos = end

    # Trailing plain text
    if pos < len(text):
        segments.append((text[pos:], {}))

    return segments


# ---------------------------------------------------------------------------
# Office Plugin Registry — allows external plugins to register document processors
# ---------------------------------------------------------------------------
_office_plugins: dict = {}

def register_office_plugin(name: str):
    """Decorator to register an office document processor plugin."""
    def decorator(func):
        _office_plugins[name] = func
        return func
    return decorator

def _call_office_plugins(plugin_type: str, *args, **kwargs) -> dict:
    """Call all registered plugins of a given type and return results dict."""
    results = {}
    for pname, plugin in _office_plugins.items():
        try:
            if callable(plugin):
                results[pname] = plugin(plugin_type, *args, **kwargs)
        except Exception as e:
            print(f"[office-plugin] '{pname}' failed on '{plugin_type}': {e}", file=sys.stderr)
    return results

def _encode_filename(filename: str) -> str:
    """Encode filename to base64 for safe SQLite3 storage (handles international chars, special chars, path traversal)."""
    if not filename:
        return filename
    safe = filename.encode('utf-8')
    encoded = _b64_mod.urlsafe_b64encode(safe).decode('ascii').rstrip('=')
    _, ext = os.path.splitext(filename)
    return encoded + ext

def _decode_filename(encoded_name: str) -> str:
    """Decode a base64-encoded filename back to original. Returns as-is if decoding fails."""
    try:
        base = os.path.splitext(encoded_name)[0]
        padding = 4 - len(base) % 4
        if padding != 4:
            base += '=' * padding
        decoded = _b64_mod.urlsafe_b64decode(base).decode('utf-8')
        return decoded + os.path.splitext(encoded_name)[1]
    except Exception:
        return encoded_name



def _read_odf(file_bytes: bytes, filename: str) -> str:
    """Read ODF files (.odt, .ods, .odp) and return structured text."""
    ext = os.path.splitext(filename)[1].lower()
    
    try:
        if ext == ".ods":
            from odf.opendocument import load
            from odf.table import Table, TableRow, TableCell
            from odf.text import P
            
            doc = load(io.BytesIO(file_bytes))
            result = []
            for table in doc.getElementsByType(Table):
                for row in table.getElementsByType(TableRow):
                    cells = []
                    for cell in row.getElementsByType(TableCell):
                        text_parts = []
                        for p in cell.getElementsByType(P):
                            try:
                                text_parts.append(str(p))
                            except Exception:
                                pass
                        cells.append(" ".join(text_parts).strip())
                    result.append(" | ".join(cells))
            return "\n".join(result) if result else "(empty spreadsheet)"
        
        elif ext == ".odt":
            from odf.opendocument import load
            from odf.text import P, H
            
            doc = load(io.BytesIO(file_bytes))
            result = []
            for elem in doc.getElementsByType(H):
                try:
                    result.append(f"## {str(elem)}")
                except Exception:
                    pass
            for elem in doc.getElementsByType(P):
                try:
                    text = str(elem).strip()
                    if text:
                        result.append(text)
                except Exception:
                    pass
            return "\n\n".join(result) if result else "(empty document)"
        
        elif ext == ".odp":
            from odf.opendocument import load
            from odf.text import P
            
            doc = load(io.BytesIO(file_bytes))
            result = []
            slide_num = 0
            for elem in doc.getElementsByType(P):
                try:
                    text = str(elem).strip()
                    if text:
                        result.append(f"Slide {slide_num + 1}: {text}")
                        slide_num += 1
                except Exception:
                    pass
            return "\n".join(result) if result else "(empty presentation)"
        
        return f"Unsupported ODF format: {ext}"
    except ImportError:
        return "Error: odfpy library not installed. Install with: pip install odfpy"
    except Exception as e:
        return f"Error reading {ext} file: {str(e)}"



# ---------------------------------------------------------------------------
# Professional document helpers
# ---------------------------------------------------------------------------
def _add_callout_box(doc, lines, colors, hex_to_rgb, fmt_mode="format"):
    """Add a professional callout box (note/tip/warning)."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    first_line = lines[0].lower() if lines else ""
    if first_line.startswith("**warning") or first_line.startswith("**alert"):
        border_color = "E74C3C"
        bg_color = "FDEDEC"
        icon = "\u26a0\ufe0f"
    elif first_line.startswith("**tip") or first_line.startswith("**pro tip"):
        border_color = "27AE60"
        bg_color = "E8F8F5"
        icon = "\U0001f4a1"
    elif first_line.startswith("**note") or first_line.startswith("**info"):
        border_color = colors.get("accent", "2E75B6")
        bg_color = colors.get("light", "D6E4F0")
        icon = "\U0001f4cc"

    table = doc.add_table(rows=1, cols=1)
    table.alignment = 1
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    tcPr.append(shading)

    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    for i, line in enumerate(lines):
        text = _format_text(line, mode=fmt_mode)
        segments = _parse_inline_md(text)
        if i == 0:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(f"{icon} ")
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = 'Calibri'
            for seg_text, fmt in segments:
                if not seg_text:
                    continue
                run = p.add_run(seg_text)
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.name = 'Calibri'
                if fmt.get('italic'):
                    run.font.italic = True
                if fmt.get('code'):
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
        else:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            for seg_text, fmt in segments:
                if not seg_text:
                    continue
                run = p.add_run(seg_text)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                if fmt.get('bold'):
                    run.font.bold = True
                if fmt.get('italic'):
                    run.font.italic = True
                if fmt.get('code'):
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)

    doc.add_paragraph()


def _add_professional_table(doc, rows, colors, hex_to_rgb, fmt_mode="format"):
    """Add a professionally styled table."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Colorful Grid Accent 1'
    table.alignment = 1

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                segments = _parse_inline_md(_format_text(cell_text, mode=fmt_mode))
                for seg_text, fmt in segments:
                    if not seg_text:
                        continue
                    run = p.add_run(seg_text)
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    if i == 0:
                        run.font.bold = True
                    if fmt.get('bold'):
                        run.font.bold = True
                    if fmt.get('italic'):
                        run.font.italic = True
                    if fmt.get('code'):
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)

    doc.add_paragraph()


def _render_content_slide(prs, lines, colors, hex_to_rgb, add_text_box, add_accent_bar, set_slide_bg, slide_num, fmt_mode="format"):
    """Render a content slide with professional layout."""
    slide = prs.slides.add_slide(prs.slide_layouts[6])
    set_slide_bg(slide, colors["bg"])

    y_pos = 0.5
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('# ') or line.startswith('## '):
            text = line.lstrip('#').strip()
            add_accent_bar(slide, 0.8, y_pos + 0.15, 0.06, 0.5, colors["accent"])
            add_text_box(slide, 1.1, y_pos, 11, 0.7, text, font_size=28, bold=True, color=colors["text"])
            y_pos += 0.9
        elif line.startswith('- ') or line.startswith('* '):
            text = _format_text(line[2:].strip(), mode=fmt_mode)
            add_text_box(slide, 1.5, y_pos, 10.5, 0.5, "\u2022 " + text, font_size=16, color=colors["text"])
            y_pos += 0.5
        elif line.startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if any(c.startswith('---') for c in cells):
                continue
            add_text_box(slide, 1.0, y_pos, 11, 0.5, "  |  ".join(cells), font_size=14, color=colors["text"])
            y_pos += 0.4
        else:
            text = _format_text(line, mode=fmt_mode)
            add_text_box(slide, 1.0, y_pos, 11.5, 0.5, text, font_size=16, color=colors["text"])
            y_pos += 0.5

        if y_pos > 6.5:
            break

__all__ = [
    "_office_plugins", "register_office_plugin", "_call_office_plugins",
    "_resolve_file_path", "_read_file_bytes", "_detect_type", "_cell_value",
    "_xls_to_xlsx", "_format_text", "_parse_inline_md", "_encode_filename",
    "_decode_filename", "_read_odf", "_add_callout_box", "_add_professional_table",
    "_render_content_slide",
]
